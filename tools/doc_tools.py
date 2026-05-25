import os
import sys

def word_to_pdf(docx_path: str, pdf_path: str = None) -> str:
    """High-quality DOCX to PDF conversion using mammoth and Qt's rendering engine."""
    try:
        import mammoth
        from PySide6.QtGui import QTextDocument, QPdfWriter, QPageLayout, QPageSize
        from PySide6.QtCore import QSizeF, QMarginsF

        if pdf_path is None:
            pdf_path = os.path.splitext(docx_path)[0] + ".pdf"

        with open(docx_path, "rb") as docx_file:
            # Convert DOCX to HTML
            result = mammoth.convert_to_html(docx_file)
            html = result.value
            
            # Basic CSS for better PDF layout
            styled_html = f"""
            <html>
            <head>
            <style>
                body {{ font-family: sans-serif; line-height: 1.5; margin: 25px; color: #333; }}
                table {{ border-collapse: collapse; width: 100%; margin: 15px 0; border: 1px solid #444; }}
                th, td {{ border: 1px solid #444; padding: 8px; text-align: left; }}
                img {{ max-width: 100%; height: auto; }}
                h1, h2, h3 {{ color: #1a2a3a; }}
                p {{ margin-bottom: 10px; }}
            </style>
            </head>
            <body>{html}</body>
            </html>
            """

        doc = QTextDocument()
        doc.setHtml(styled_html)

        writer = QPdfWriter(pdf_path)
        writer.setPageSize(QPageSize(QPageSize.PageSizeId.A4))
        writer.setPageMargins(QMarginsF(15, 15, 15, 15), QPageLayout.Unit.Millimeter)
        
        # Render the document to PDF
        doc.print_(writer)
        return pdf_path

    except Exception as e:
        print(f"Enhanced Word to PDF failed: {e}. Using fallback...")
        # Fallback to fpdf method if mammoth/Qt fails
        from docx import Document
        from fpdf import FPDF

        if pdf_path is None:
            pdf_path = os.path.splitext(docx_path)[0] + ".pdf"

        doc = Document(docx_path)
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Helvetica", "", 12)

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                try:
                    pdf.multi_cell(0, 10, text)
                except:
                    pdf.multi_cell(0, 10, text.encode("latin-1", "replace").decode("latin-1"))
            else:
                pdf.ln(5)

        pdf.output(pdf_path)
        return pdf_path


def pdf_to_word(pdf_path: str, docx_path: str = None) -> str:
    """High-quality PDF to DOCX conversion using layout analysis."""
    try:
        from pdf2docx import Converter
        if docx_path is None:
            docx_path = os.path.splitext(pdf_path)[0] + ".docx"
        
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
        return docx_path
    except Exception as e:
        print(f"Advanced PDF to Word failed: {e}. Using text extraction fallback...")
        from docx import Document
        from pypdf import PdfReader

        if docx_path is None:
            docx_path = os.path.splitext(pdf_path)[0] + ".docx"

        doc = Document()
        reader = PdfReader(pdf_path)
        for page in reader.pages:
            text = page.extract_text()
            if text:
                doc.add_paragraph(text)
        doc.save(docx_path)
        return docx_path


def image_to_pdf(image_paths: list, pdf_path: str = "output.pdf") -> str:
    import img2pdf
    with open(pdf_path, "wb") as f:
        f.write(img2pdf.convert(image_paths))
    return pdf_path


def split_pdf(pdf_path: str, output_dir: str = None) -> list:
    import pikepdf

    if output_dir is None:
        output_dir = os.path.dirname(pdf_path) or "."
    os.makedirs(output_dir, exist_ok=True)

    base = os.path.splitext(os.path.basename(pdf_path))[0]
    paths = []

    with pikepdf.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages):
            out = pikepdf.Pdf.new()
            out.pages.append(page)
            out_path = os.path.join(output_dir, f"{base}_page_{i+1}.pdf")
            out.save(out_path)
            out.close()
            paths.append(out_path)

    return paths


def merge_pdfs(pdf_list: list, output_path: str = "merged.pdf") -> str:
    from tools.pdf import merge_documents
    merge_documents(pdf_list, output_path)
    return output_path


def text_to_pdf(text: str, pdf_path: str = "output.pdf") -> str:
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Helvetica", "", 12)

    for line in text.splitlines():
        try:
            pdf.multi_cell(0, 10, line)
        except:
            pdf.multi_cell(0, 10, line.encode("latin-1", "replace").decode("latin-1"))
    pdf.output(pdf_path)
    return pdf_path
